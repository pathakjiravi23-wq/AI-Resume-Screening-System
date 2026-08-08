from pypdf import PdfReader
from docx import Document


def pdf_file(path) -> str:
    docu = PdfReader(path)
    text = []
    for page in docu.pages:
        text.append(page.extract_text())
    return "\n".join(text)


def docx_file(path) -> str:
    docu = Document(path)
    text = []
    for para in docu.paragraphs:
        if para.text.strip():
            text.append(para.text)
    # Extract tables
    for tables in docu.tables:
        for row in tables.rows:
            row_cells = "|".join(cell.text.strip() for cell in row.cells)
            text.append(row_cells)
    return "\n".join(text)


def text_file(path) -> str:
    with open(path, "r", encoding="UTF-8") as txtfile:
        return txtfile.read()
